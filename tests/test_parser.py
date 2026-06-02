"""
DocumentParser 单元测试

覆盖：
- 文本解析（TXT）
- DOCX 解析
- PDF 解析（mock）
- 条款切分逻辑
- 错误处理（不支持格式、损坏文件）
- parse_file vs parse_bytes 一致性
"""

import os
import tempfile

import pytest

from src.core.exceptions import FileCorruptedError, UnsupportedFormatError
from src.parsing.parser import Clause, DocumentParser


@pytest.fixture
def parser():
    return DocumentParser()


# ============================================
# 文本文件解析
# ============================================


class TestParseTxt:
    """TXT 文件解析测试"""

    def test_parse_txt_bytes_utf8(self, parser):
        """UTF-8 编码的 TXT 文件"""
        content = "第一条 违约责任\n如乙方违约，应承担违约责任。"
        result = parser.parse_bytes(content.encode("utf-8"), "test.txt")

        assert result.file_type == "txt"
        assert "违约责任" in result.full_text
        assert result.original_filename == "test.txt"

    def test_parse_txt_bytes_gbk(self, parser):
        """GBK 编码的 TXT 文件（errors='replace' 兜底）"""
        content = "合同条款"
        gbk_bytes = content.encode("gbk")
        result = parser.parse_bytes(gbk_bytes, "test.txt")

        assert result.file_type == "txt"
        assert len(result.full_text) > 0

    def test_parse_txt_file(self, parser):
        """从文件路径解析 TXT"""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("第一条 甲方义务\n甲方应按时付款。")
            tmp_path = f.name

        try:
            result = parser.parse_file(tmp_path)
            assert result.file_type == "txt"
            assert "甲方义务" in result.full_text
            assert result.metadata["extension"] == ".txt"
        finally:
            os.unlink(tmp_path)

    def test_parse_txt_empty(self, parser):
        """空 TXT 文件"""
        result = parser.parse_bytes(b"", "empty.txt")
        assert result.file_type == "txt"
        assert result.full_text == ""

    def test_parse_txt_metadata(self, parser):
        """验证 metadata 包含文件大小"""
        content = "测试内容"
        result = parser.parse_bytes(content.encode("utf-8"), "test.txt")
        assert result.metadata["file_size"] == len(content.encode("utf-8"))
        assert result.metadata["extension"] == ".txt"


# ============================================
# DOCX 解析
# ============================================


class TestParseDocx:
    """DOCX 文件解析测试"""

    def test_parse_docx_bytes(self, parser):
        """解析 DOCX 字节流（需要 python-docx）"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")

        doc = Document()
        doc.add_paragraph("第一条 违约责任")
        doc.add_paragraph("如乙方违约，应承担违约责任。")

        import io

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = parser.parse_bytes(buf.read(), "test.docx")
        assert result.file_type == "docx"
        assert "违约责任" in result.full_text

    def test_parse_docx_empty_paragraphs_skipped(self, parser):
        """DOCX 中空段落应被跳过"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")

        doc = Document()
        doc.add_paragraph("有效内容")
        doc.add_paragraph("")
        doc.add_paragraph("另一段有效内容")

        import io

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = parser.parse_bytes(buf.read(), "test.docx")
        lines = result.full_text.split("\n")
        assert all(line.strip() for line in lines)


# ============================================
# 错误处理
# ============================================


class TestParserErrors:
    """解析器错误处理测试"""

    def test_unsupported_format_xlsx(self, parser):
        """不支持的文件格式应抛出 UnsupportedFormatError"""
        with pytest.raises(UnsupportedFormatError) as exc_info:
            parser.parse_bytes(b"test", "test.xlsx")
        assert ".xlsx" in str(exc_info.value)

    def test_unsupported_format_exe(self, parser):
        """exe 文件应抛出 UnsupportedFormatError"""
        with pytest.raises(UnsupportedFormatError):
            parser.parse_bytes(b"MZ", "malware.exe")

    def test_file_not_found(self, parser):
        """文件不存在应抛出 FileNotFoundError"""
        with pytest.raises(FileNotFoundError):
            parser.parse_file("/nonexistent/path/file.txt")

    def test_corrupted_pdf(self, parser):
        """损坏的 PDF 应抛出 FileCorruptedError"""
        with pytest.raises(FileCorruptedError):
            parser.parse_bytes(b"not a real pdf content", "test.pdf")

    def test_corrupted_docx(self, parser):
        """损坏的 DOCX 应抛出 FileCorruptedError"""
        with pytest.raises(FileCorruptedError):
            parser.parse_bytes(b"PK\x03\x04corrupted", "test.docx")


# ============================================
# 条款切分
# ============================================


class TestClauseSplitting:
    """条款切分逻辑测试"""

    def test_split_by_article_number(self, parser):
        """按「第X条」切分"""
        text = "第一条 甲方义务\n甲方应按时付款。\n\n第二条 乙方义务\n乙方应按时交货。"
        result = parser.parse_bytes(text.encode("utf-8"), "test.txt")

        assert len(result.clauses) >= 2
        titles = [c.title for c in result.clauses if c.title]
        assert any("第一条" in t for t in titles)
        assert any("第二条" in t for t in titles)

    def test_split_by_chinese_number(self, parser):
        """按中文数字序号切分（一、二、三）"""
        text = "一、合同目的\n本合同旨在明确双方权利义务。\n\n二、合作内容\n双方合作开发软件。"
        result = parser.parse_bytes(text.encode("utf-8"), "test.txt")

        assert len(result.clauses) >= 2

    def test_split_by_keyword_title(self, parser):
        """按常见条款标题关键词切分"""
        text = "违约责任\n如一方违约，应赔偿对方损失。\n\n争议解决\n双方应友好协商解决争议。"
        result = parser.parse_bytes(text.encode("utf-8"), "test.txt")

        assert len(result.clauses) >= 2

    def test_no_structure_fallback_to_paragraph(self, parser):
        """无结构化标题时，按段落切分"""
        text = "这是第一段内容。\n\n这是第二段内容。\n\n这是第三段内容。"
        result = parser.parse_bytes(text.encode("utf-8"), "test.txt")

        assert len(result.clauses) >= 1

    def test_single_line_document(self, parser):
        """单行文档作为单个条款"""
        text = "这是一份简单的合同文本。"
        result = parser.parse_bytes(text.encode("utf-8"), "test.txt")

        assert len(result.clauses) == 1
        assert result.clauses[0].id == 1

    def test_clause_id_sequential(self, parser):
        """条款 ID 应连续递增"""
        text = "第一条 甲方义务\n甲方应付款。\n第二条 乙方义务\n乙方应交货。\n第三条 违约责任\n违约方赔偿。"
        result = parser.parse_bytes(text.encode("utf-8"), "test.txt")

        ids = [c.id for c in result.clauses]
        assert ids == list(range(1, len(result.clauses) + 1))

    def test_clause_content_preserved(self, parser):
        """条款内容应完整保留"""
        text = "第一条 保密条款\n双方应对本合同内容保密，未经对方书面同意，不得向第三方披露。"
        result = parser.parse_bytes(text.encode("utf-8"), "test.txt")

        assert any("保密" in c.content for c in result.clauses)


# ============================================
# 条款类型检测
# ============================================


class TestClauseTypeDetection:
    """条款类型检测测试"""

    def test_detect_breach_clause(self, parser):
        """检测违约责任条款"""
        clause = Clause(id=1, content="如乙方违约，应支付违约金10万元。", title="违约责任")
        clause_type = parser.detect_clause_type(clause)
        assert clause_type == "违约责任"

    def test_detect_dispute_clause(self, parser):
        """检测争议解决条款"""
        clause = Clause(id=1, content="双方应友好协商解决争议，协商不成的，提交仲裁委员会仲裁。", title="争议解决")
        clause_type = parser.detect_clause_type(clause)
        assert clause_type == "争议解决"

    def test_detect_confidentiality_clause(self, parser):
        """检测保密条款"""
        clause = Clause(id=1, content="双方应对合同内容及商业秘密严格保密。", title="保密条款")
        clause_type = parser.detect_clause_type(clause)
        assert clause_type == "保密条款"

    def test_detect_unknown_clause(self, parser):
        """无法识别的条款类型返回 None"""
        clause = Clause(id=1, content="今天天气很好。", title="无关内容")
        clause_type = parser.detect_clause_type(clause)
        assert clause_type is None


# ============================================
# parse_file 与 parse_bytes 一致性
# ============================================


class TestParseConsistency:
    """parse_file 和 parse_bytes 结果一致性"""

    def test_txt_consistency(self, parser):
        """TXT 文件：parse_file 和 parse_bytes 应产生相同文本"""
        content = "第一条 违约责任\n如乙方违约，应承担违约责任。"

        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write(content)
            tmp_path = f.name

        try:
            file_result = parser.parse_file(tmp_path)
            bytes_result = parser.parse_bytes(content.encode("utf-8"), os.path.basename(tmp_path))

            assert file_result.full_text == bytes_result.full_text
            assert len(file_result.clauses) == len(bytes_result.clauses)
        finally:
            os.unlink(tmp_path)
