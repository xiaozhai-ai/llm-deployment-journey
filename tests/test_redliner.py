"""
Redliner 单元测试

覆盖：
- HTML 差异对比生成
- XSS 防护（clause_title / risk_name 转义）
- DOCX 生成（含声明文本）
- JSON 解析容错（测试真实 _extract_json 方法）
- 空修订处理
"""

import io
import json
from unittest.mock import MagicMock

import pytest

from src.output.redliner import ClauseRevision

# ============================================
# 测试用数据
# ============================================


def _make_revision(**kwargs):
    defaults = dict(
        clause_id=1,
        clause_title="第一条 付款方式",
        original_text="甲方应在合同签订后30日内支付全款。",
        revised_text="甲方应在合同签订后30日内支付全款，逾期按日万分之五支付违约金。",
        revision_type="modify",
        risk_id="R1",
        risk_name="违约责任缺失",
        explanation="补充违约金条款",
        html_diff="",
    )
    defaults.update(kwargs)
    return ClauseRevision(**defaults)


# ============================================
# HTML 差异对比
# ============================================


class TestHtmlDiff:
    def test_diff_shows_additions(self, redliner):
        html = redliner._generate_html_diff("原条款", "原条款\n新增内容")
        assert "diff-added" in html
        assert "新增内容" in html

    def test_diff_shows_removals(self, redliner):
        html = redliner._generate_html_diff("原条款\n删除内容", "原条款")
        assert "diff-removed" in html
        assert "删除内容" in html

    def test_diff_no_changes(self, redliner):
        html = redliner._generate_html_diff("相同文本", "相同文本")
        assert "无差异" in html

    def test_diff_escapes_html(self, redliner):
        html = redliner._generate_html_diff("原<条款>", "新<条款>")
        assert "<条款>" not in html
        assert "&lt;条款&gt;" in html


# ============================================
# 完整 HTML 对比报告
# ============================================


class TestFullHtmlDiff:
    def test_no_revisions(self, redliner):
        html = redliner._generate_full_html_diff("原文", [])
        assert "无需修订" in html

    def test_single_revision(self, redliner):
        rev = _make_revision()
        html = redliner._generate_full_html_diff("原文", [rev])
        assert "共 1 处修订建议" in html
        assert "违约责任缺失" in html

    def test_xss_in_risk_name(self, redliner):
        rev = _make_revision(risk_name='<script>alert("xss")</script>')
        html = redliner._generate_full_html_diff("原文", [rev])
        assert "<script>" not in html
        assert "&lt;script&gt;" in html

    def test_xss_in_clause_title(self, redliner):
        rev = _make_revision(clause_title="<img onerror=alert(1)>")
        html = redliner._generate_full_html_diff("原文", [rev])
        assert "<img" not in html
        assert "&lt;img" in html

    def test_xss_in_original_text(self, redliner):
        rev = _make_revision(original_text="<b>加粗</b>")
        html = redliner._generate_full_html_diff("原文", [rev])
        assert "<b>" not in html

    def test_multiple_revisions(self, redliner):
        revisions = [_make_revision(risk_id=f"R{i}", risk_name=f"风险{i}") for i in range(3)]
        html = redliner._generate_full_html_diff("原文", revisions)
        assert "共 3 处修订建议" in html


# ============================================
# DOCX 生成
# ============================================


class TestDocxGeneration:
    def test_generates_valid_docx(self, redliner):
        rev = _make_revision()
        docx_bytes = redliner.generate_docx_with_revisions("原文", [rev])
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
        # DOCX 文件头 (PK zip)
        assert docx_bytes[:2] == b"PK"

    def test_docx_contains_disclaimer(self, redliner):
        rev = _make_revision()
        docx_bytes = redliner.generate_docx_with_revisions("原文", [rev])
        # 解析 DOCX 验证免责声明
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "非 Word 原生修订标记" in full_text
        assert "不构成正式法律意见" in full_text

    def test_docx_contains_revision_details(self, redliner):
        rev = _make_revision()
        docx_bytes = redliner.generate_docx_with_revisions("原文", [rev])
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "原条款" in full_text
        assert "修订后" in full_text
        assert "违约责任缺失" in full_text

    def test_docx_empty_revisions(self, redliner):
        docx_bytes = redliner.generate_docx_with_revisions("原文", [])
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0


# ============================================
# JSON 解析容错（测试真实 _extract_json 方法）
# ============================================


class TestJsonParsing:
    def test_valid_json_response(self, redliner):
        """LLM 返回标准 JSON 时应正常解析"""
        response = '{"revised_text": "修订后文本", "explanation": "理由"}'
        result = redliner._extract_json(response)
        assert result is not None
        data = json.loads(result)
        assert data["revised_text"] == "修订后文本"

    def test_json_with_surrounding_text(self, redliner):
        """LLM 在 JSON 前后附加说明文字时应正确提取"""
        response = '好的，以下是修订结果：\n{"revised_text": "新条款", "explanation": "理由"}\n以上是修订。'
        result = redliner._extract_json(response)
        assert result is not None
        data = json.loads(result)
        assert data["revised_text"] == "新条款"

    def test_json_in_markdown_fence(self, redliner):
        """Markdown 代码块中的 JSON 应正确提取"""
        response = '```json\n{"revised_text": "条款内容", "explanation": "理由"}\n```'
        result = redliner._extract_json(response)
        assert result is not None
        data = json.loads(result)
        assert data["revised_text"] == "条款内容"

    def test_json_in_plain_code_fence(self, redliner):
        """普通代码块中的 JSON 应正确提取"""
        response = '```\n{"revised_text": "条款内容"}\n```'
        result = redliner._extract_json(response)
        assert result is not None
        data = json.loads(result)
        assert data["revised_text"] == "条款内容"

    def test_nested_json_structure(self, redliner):
        """嵌套 JSON 结构应正确提取"""
        response = '结果：\n{"revised_text": "新条款", "details": {"reason": "合规"}}\n以上。'
        result = redliner._extract_json(response)
        assert result is not None
        data = json.loads(result)
        assert data["revised_text"] == "新条款"
        assert data["details"]["reason"] == "合规"

    def test_no_json_returns_none(self, redliner):
        """无 JSON 内容时应返回 None"""
        response = "这是一个纯文本响应，没有 JSON"
        result = redliner._extract_json(response)
        assert result is None

    def test_empty_string_returns_none(self, redliner):
        """空字符串应返回 None"""
        result = redliner._extract_json("")
        assert result is None

    def test_malformed_json_in_code_fence_returns_none(self, redliner):
        """代码块中畸形 JSON 应返回 None"""
        response = '```json\n{"revised_text": "未闭合\n```'
        result = redliner._extract_json(response)
        assert result is None


# ============================================
# 修订生成（异步）
# ============================================


class TestGenerateRevisions:
    @pytest.mark.asyncio
    async def test_no_actionable_risks(self, redliner):
        """无建议的风险不生成修订"""
        risks = [MagicMock(suggestion=None, clause_content_preview="文本")]
        doc = await redliner.generate_revisions("原文", risks)
        assert len(doc.revisions) == 0

    @pytest.mark.asyncio
    async def test_short_preview_skipped(self, redliner):
        """过短的条款预览跳过"""
        risks = [MagicMock(suggestion="建议", clause_content_preview="短")]
        doc = await redliner.generate_revisions("原文", risks)
        assert len(doc.revisions) == 0

    @pytest.mark.asyncio
    async def test_max_five_revisions(self, redliner):
        """最多生成 5 个修订"""
        risks = []
        for i in range(10):
            r = MagicMock()
            r.suggestion = f"建议{i}"
            r.clause_content_preview = "甲" * 50
            r.clause_position = f"第{i}条"
            r.risk_level = "high"
            risks.append(r)

        doc = await redliner.generate_revisions("原文", risks)
        assert len(doc.revisions) <= 5
