"""
修订生成器模块 (Redliner)
- LLM 生成修订后条款文本
- HTML 差异对比视图（增绿/删红/改黄）
- DOCX Track Changes 导出
- 增强版：统一异常处理 + 日志记录
"""

import difflib
import io
import re
from dataclasses import dataclass, field

from src.exceptions import DOCXGenerationError, RevisionError
from src.logger import logger_manager


@dataclass
class ClauseRevision:
    """条款修订"""

    clause_id: int
    clause_title: str | None
    original_text: str
    revised_text: str
    revision_type: str  # modify / delete / add
    risk_id: str  # 关联的风险 ID
    risk_name: str
    explanation: str  # 修改说明
    html_diff: str = ""  # HTML 差异对比


@dataclass
class RevisionDocument:
    """修订文档"""

    filename: str
    revisions: list[ClauseRevision] = field(default_factory=list)
    html_full_diff: str = ""  # 完整 HTML 对比
    docx_bytes: bytes | None = None  # DOCX 文件字节流


class Redliner:
    """修订生成器"""

    def __init__(self, llm_client=None):
        """
        初始化修订生成器

        Args:
            llm_client: LLM 客户端实例
        """
        self.llm_client = llm_client

    def set_llm_client(self, llm_client):
        """设置 LLM 客户端"""
        self.llm_client = llm_client

    async def generate_revisions(
        self,
        document_text: str,
        risks: list,  # List[RiskItem]
        playbook=None,
    ) -> RevisionDocument:
        """
        为识别的风险生成修订建议

        Args:
            document_text: 原始文档文本
            risks: 风险列表
            playbook: 审查策略（影响修订倾向）

        Returns:
            RevisionDocument: 修订文档
        """
        revisions = []

        # 仅对有具体建议的风险生成修订，按风险等级排序，最多 5 个
        actionable_risks = [r for r in risks if r.suggestion and r.clause_content_preview]
        _LEVEL_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3}
        actionable_risks.sort(key=lambda r: _LEVEL_ORDER.get(getattr(r, "risk_level", "low"), 3))
        actionable_risks = actionable_risks[:5]

        for risk in actionable_risks:
            revision = await self._generate_clause_revision(document_text, risk, playbook)
            if revision:
                revisions.append(revision)

        # 生成 HTML 完整对比
        html_full_diff = self._generate_full_html_diff(document_text, revisions)

        return RevisionDocument(filename="revised_document", revisions=revisions, html_full_diff=html_full_diff)

    async def _generate_clause_revision(
        self,
        document_text: str,
        risk,  # RiskItem
        playbook=None,
    ) -> ClauseRevision | None:
        """为单个风险生成条款修订"""
        original_text = risk.clause_content_preview or ""
        if len(original_text) < 10:
            return None

        # 使用 LLM 生成修订文本
        if self.llm_client:
            revised_text, explanation = await self._llm_generate_revision(original_text, risk, playbook)
        else:
            # 无 LLM 时基于建议生成简单修订
            revised_text = risk.suggestion
            explanation = "基于规则建议生成"

        if not revised_text:
            return None

        # 生成 HTML 差异对比
        html_diff = self._generate_html_diff(original_text, revised_text)

        return ClauseRevision(
            clause_id=0,
            clause_title=risk.clause_position,
            original_text=original_text,
            revised_text=revised_text,
            revision_type="modify",
            risk_id=risk.id,
            risk_name=risk.name,
            explanation=explanation,
            html_diff=html_diff,
        )

    async def _llm_generate_revision(self, original_text: str, risk, playbook=None) -> tuple[str, str]:
        """使用 LLM 生成修订文本"""
        role_context = ""
        if playbook:
            role_map = {"party_a": "你代表甲方立场", "party_b": "你代表乙方立场", "neutral": "你保持中立立场"}
            role_context = role_map.get(playbook.role, "")

        prompt = f"""{role_context}修订以下条款，返回 JSON：
{{"revised_text":"修订后完整条款","explanation":"修改理由(≤80字)"}}

风险：{risk.description}
依据：{risk.legal_basis or "无"}
建议：{risk.suggestion}
原条款：{original_text}"""

        try:
            response = await self.llm_client.chat_completion(
                prompt, system_prompt="法律文书修订助手。仅返回 JSON。", temperature=0.3, max_tokens=1000
            )

            import json

            json_str = None
            # 策略1：匹配包含 revised_text 的完整 JSON 对象
            m = re.search(r'\{[\s\S]*?"revised_text"[\s\S]*?\}', response)
            if m:
                json_str = m.group()
            # 策略2：从 markdown 代码块中提取
            if not json_str:
                m = re.search(r"```(?:json)?\s*(\{[\s\S]*?\})\s*```", response)
                if m:
                    json_str = m.group(1)
            # 策略3：匹配任意 JSON 对象
            if not json_str:
                m = re.search(r"\{[\s\S]*?\}", response)
                if m:
                    json_str = m.group()

            if json_str:
                try:
                    data = json.loads(json_str)
                    return data.get("revised_text", ""), data.get("explanation", "")
                except json.JSONDecodeError:
                    logger_manager.warning(f"LLM 修订 JSON 解析失败: {json_str[:200]}")

            logger_manager.warning(f"LLM 修订响应格式异常: {response[:200]}")
            return "", "LLM 响应格式异常"

        except RevisionError:
            raise
        except Exception as e:
            logger_manager.error(f"LLM 修订生成失败: {e}")
            raise RevisionError(f"LLM 修订生成失败: {e}", error_code="LLM_REVISION_ERROR") from e

    def _generate_html_diff(self, original: str, revised: str) -> str:
        """
        生成 HTML 差异对比

        Args:
            original: 原始文本
            revised: 修订后文本

        Returns:
            HTML 差异对比字符串
        """
        # 按行分割
        orig_lines = original.split("\n")
        rev_lines = revised.split("\n")

        diff = difflib.unified_diff(orig_lines, rev_lines, lineterm="", n=0)

        html_parts = []
        html_parts.append('<div class="diff-container">')
        html_parts.append("<style>")
        html_parts.append("""
            .diff-added { background-color: #d4edda; color: #155724; }
            .diff-removed { background-color: #f8d7da; color: #721c24; text-decoration: line-through; }
            .diff-header { color: #6c757d; font-size: 0.9em; margin: 5px 0; }
            .diff-content { font-family: "Microsoft YaHei", sans-serif; line-height: 1.6; padding: 10px; border-radius: 4px; }
        """)
        html_parts.append("</style>")
        html_parts.append('<div class="diff-content">')

        has_changes = False
        for line in diff:
            if line.startswith("---") or line.startswith("+++"):
                continue
            if line.startswith("@@"):
                html_parts.append('<div class="diff-header">条款修订</div>')
                continue
            if line.startswith("+"):
                has_changes = True
                html_parts.append(f'<span class="diff-added">{self._escape_html(line[1:])}</span><br>')
            elif line.startswith("-"):
                has_changes = True
                html_parts.append(f'<span class="diff-removed">{self._escape_html(line[1:])}</span><br>')
            else:
                html_parts.append(f"{self._escape_html(line)}<br>")

        if not has_changes:
            html_parts.append('<div class="diff-header">无差异</div>')

        html_parts.append("</div></div>")
        return "\n".join(html_parts)

    def _generate_full_html_diff(self, full_text: str, revisions: list[ClauseRevision]) -> str:
        """
        生成完整文档的 HTML 对比报告

        Args:
            full_text: 完整原始文本
            revisions: 修订列表

        Returns:
            完整 HTML 报告
        """
        html = []
        html.append('<div class="full-diff-report">')
        html.append("<style>")
        html.append("""
            .full-diff-report { font-family: "Microsoft YaHei", sans-serif; }
            .revision-block { margin: 20px 0; padding: 15px; border: 1px solid #dee2e6; border-radius: 8px; }
            .revision-block h4 { margin: 0 0 10px 0; color: #495057; }
            .revision-block .risk-tag {
                display: inline-block; padding: 2px 8px; border-radius: 4px;
                font-size: 0.85em; margin-right: 8px;
            }
            .risk-high { background-color: #f8d7da; color: #721c24; }
            .risk-medium { background-color: #fff3cd; color: #856404; }
            .risk-low { background-color: #d4edda; color: #155724; }
            .original-text {
                background-color: #f8f9fa; padding: 10px; border-radius: 4px;
                margin: 10px 0; border-left: 3px solid #dc3545;
            }
            .revised-text {
                background-color: #f8f9fa; padding: 10px; border-radius: 4px;
                margin: 10px 0; border-left: 3px solid #28a745;
            }
            .explanation { color: #6c757d; font-size: 0.9em; margin-top: 10px; }
        """)
        html.append("</style>")

        if not revisions:
            html.append("<p>无需修订或无可操作的修订建议。</p>")
        else:
            html.append(f"<h3>共 {len(revisions)} 处修订建议</h3>")

            for i, rev in enumerate(revisions, 1):
                escaped_title = self._escape_html(rev.clause_title) if rev.clause_title else ""
                html.append(f"""
                <div class="revision-block">
                    <h4>📝 修订 {i}: {self._escape_html(rev.risk_name)}</h4>
                    {escaped_title}
                    <div class="original-text">
                        <strong>📄 原条款:</strong><br>
                        {self._escape_html(rev.original_text)}
                    </div>
                    <div class="revised-text">
                        <strong>✅ 修订后:</strong><br>
                        {self._escape_html(rev.revised_text)}
                    </div>
                    <div class="explanation">
                        💡 {self._escape_html(rev.explanation)}
                    </div>
                </div>
                """)

        html.append("</div>")
        return "\n".join(html)

    def generate_docx_with_revisions(
        self, original_text: str, revisions: list[ClauseRevision], filename: str = "revised_document"
    ) -> bytes:
        """
        生成带修订痕迹的 DOCX 文件

        Args:
            original_text: 原始文本
            revisions: 修订列表
            filename: 文件名

        Returns:
            DOCX 文件字节流
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt, RGBColor
        except ImportError as e:
            logger_manager.error(f"python-docx 未安装: {e}")
            raise DOCXGenerationError("请安装 python-docx: pip install python-docx") from e

        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(11)

        # ── 首页：免责声明（加粗红色边框框） ──
        doc.add_heading("法律文件修订报告", level=1)

        # 免责声明放在最前面，确保用户第一时间看到
        disclaimer = doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = disclaimer.add_run("⚠️ 重 要 免 责 声 明 ⚠️")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(204, 0, 0)

        disclaimer_body = doc.add_paragraph()
        disclaimer_body.alignment = WD_ALIGN_PARAGRAPH.LEFT
        disclaimer_lines = [
            "1. 本报告由 AI 辅助生成，仅供参考，不构成正式法律意见。",
            "2. 所有修订建议须经专业律师审核确认后方可采用。",
            '3. 本文档使用颜色标注方式展示修订（非 Word 原生修订标记），无法使用"接受/拒绝修订"功能。',
            "4. 生成方不对因直接使用本报告内容而产生的任何损失承担责任。",
        ]
        for line in disclaimer_lines:
            run = disclaimer_body.add_run(line + "\n")
            run.font.color.rgb = RGBColor(204, 0, 0)
            run.font.size = Pt(10)

        # ── 颜色说明图例 ──
        doc.add_heading("颜色说明", level=2)
        legend_items = [
            ("🔴 红色标注", "原条款文本（需要修改的内容）", RGBColor(180, 0, 0)),
            ("🟢 绿色标注", "修订后条款文本（建议替换为）", RGBColor(0, 128, 0)),
            ("🟡 黄色背景", "差异高亮区域", RGBColor(128, 100, 0)),
        ]
        for icon_label, desc, color in legend_items:
            p = doc.add_paragraph()
            run = p.add_run(icon_label + "：")
            run.bold = True
            run.font.color.rgb = color
            p.add_run(desc)

        doc.add_paragraph("——" * 20)

        # ── 修订摘要 ──
        doc.add_heading("修订摘要", level=2)
        if revisions:
            doc.add_paragraph(f"共识别 {len(revisions)} 处需修订条款：")
            for i, rev in enumerate(revisions, 1):
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{i}. {rev.risk_name}")
                run.bold = True
                p.add_run(f" — {rev.explanation}")
        else:
            doc.add_paragraph("未发现需要修订的条款。")

        # ── 修订详情 ──
        if revisions:
            doc.add_heading("修订详情", level=2)

            for i, rev in enumerate(revisions, 1):
                if i > 1:
                    doc.add_page_break()

                doc.add_heading(f"修订 {i}/{len(revisions)}：{rev.risk_name}", level=3)

                # 关联信息
                if rev.clause_title:
                    p = doc.add_paragraph()
                    run = p.add_run("条款位置：")
                    run.bold = True
                    p.add_run(rev.clause_title)

                if rev.risk_id:
                    p = doc.add_paragraph()
                    run = p.add_run("风险编号：")
                    run.bold = True
                    p.add_run(rev.risk_id)

                doc.add_paragraph()

                # 原条款
                p = doc.add_paragraph()
                run = p.add_run("📄 原条款：")
                run.bold = True
                run.font.color.rgb = RGBColor(180, 0, 0)
                run.font.size = Pt(11)
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.3)
                p2.add_run(rev.original_text)

                doc.add_paragraph()

                # 修订后
                p = doc.add_paragraph()
                run = p.add_run("✅ 修订建议：")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)
                run.font.size = Pt(11)
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.3)
                p2.add_run(rev.revised_text)

                doc.add_paragraph()

                # 修改理由
                p = doc.add_paragraph()
                run = p.add_run("💡 修改理由：")
                run.bold = True
                p.add_run(rev.explanation)

        # ── 页脚免责声明 ──
        doc.add_paragraph()
        doc.add_paragraph("——" * 20)
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            "本报告由「合同卫士」AI 审查系统自动生成 | 生成时间："
            + __import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M")
        )
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.size = Pt(9)

        # 保存到字节流
        try:
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            logger_manager.error(f"DOCX 保存失败: {e}")
            raise DOCXGenerationError(f"DOCX 文件保存失败: {e}") from e

    @staticmethod
    def _escape_html(text: str) -> str:
        """HTML 转义"""
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
