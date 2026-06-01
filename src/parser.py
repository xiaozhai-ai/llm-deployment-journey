"""
文件解析模块
支持 PDF、Word (.docx)、TXT 格式文件的解析与结构化
增强版：统一异常处理 + 日志记录
"""

import os
from dataclasses import dataclass

from src.exceptions import FileCorruptedError, UnsupportedFormatError
from src.logger import logger_manager


@dataclass
class Clause:
    """条款数据结构"""

    id: int
    content: str
    title: str | None = None
    start_pos: int = 0
    end_pos: int = 0
    clause_type: str | None = None  # 条款类型（如违约责任、争议解决等）


@dataclass
class ParsedDocument:
    """解析后的文档结构"""

    original_filename: str
    file_type: str  # pdf, docx, txt
    full_text: str
    clauses: list[Clause]
    metadata: dict = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentParser:
    """文档解析器"""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}

    def __init__(self):
        pass

    def parse_file(self, file_path: str) -> ParsedDocument:
        """
        解析文件，返回结构化文档

        Args:
            file_path: 文件路径

        Returns:
            ParsedDocument: 解析后的文档结构

        Raises:
            UnsupportedFormatError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedFormatError(ext)

        if ext == ".pdf":
            full_text = self._parse_pdf(file_path)
        elif ext == ".docx":
            full_text = self._parse_docx(file_path)
        elif ext == ".doc":
            full_text = self._parse_doc(file_path)
        elif ext == ".txt":
            full_text = self._parse_txt(file_path)
        else:
            raise UnsupportedFormatError(ext)

        clauses = self._split_clauses(full_text)

        return ParsedDocument(
            original_filename=os.path.basename(file_path),
            file_type=ext[1:],
            full_text=full_text,
            clauses=clauses,
            metadata={"file_size": os.path.getsize(file_path), "extension": ext},
        )

    def parse_bytes(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """
        从字节流解析文件（用于 Web 上传）

        Args:
            file_bytes: 文件字节流
            filename: 文件名

        Returns:
            ParsedDocument: 解析后的文档结构
        """
        ext = os.path.splitext(filename)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedFormatError(ext)

        if ext == ".pdf":
            full_text = self._parse_pdf_bytes(file_bytes)
        elif ext == ".docx":
            full_text = self._parse_docx_bytes(file_bytes)
        elif ext == ".doc":
            full_text = self._parse_doc_bytes(file_bytes)
        elif ext == ".txt":
            full_text = self._parse_txt_bytes(file_bytes)
        else:
            raise UnsupportedFormatError(ext)

        clauses = self._split_clauses(full_text)

        return ParsedDocument(
            original_filename=filename,
            file_type=ext[1:],
            full_text=full_text,
            clauses=clauses,
            metadata={"file_size": len(file_bytes), "extension": ext},
        )

    def _parse_pdf(self, file_path: str) -> str:
        """解析 PDF 文件"""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            logger_manager.debug(f"成功解析 PDF 文件: {file_path}")
            return "\n".join(text_parts)
        except Exception as e:
            logger_manager.error(f"PDF 解析失败: {file_path}: {e}")
            raise FileCorruptedError(f"PDF 文件解析失败: {e}") from e

    def _parse_pdf_bytes(self, file_bytes: bytes) -> str:
        """从字节流解析 PDF"""
        try:
            import io

            import pdfplumber

            text_parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            logger_manager.debug("成功解析 PDF 字节流")
            return "\n".join(text_parts)
        except Exception as e:
            logger_manager.error(f"PDF 字节流解析失败: {e}")
            raise FileCorruptedError(f"PDF 字节流解析失败: {e}") from e

    def _parse_docx(self, file_path: str) -> str:
        """解析 Word 文件"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            logger_manager.debug(f"成功解析 DOCX 文件: {file_path}")
            return "\n".join(text_parts)
        except Exception as e:
            logger_manager.error(f"DOCX 解析失败: {file_path}: {e}")
            raise FileCorruptedError(f"Word 文件解析失败: {e}") from e

    def _parse_docx_bytes(self, file_bytes: bytes) -> str:
        """从字节流解析 Word"""
        try:
            import io

            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            logger_manager.debug("成功解析 DOCX 字节流")
            return "\n".join(text_parts)
        except Exception as e:
            logger_manager.error(f"DOCX 字节流解析失败: {e}")
            raise FileCorruptedError(f"Word 字节流解析失败: {e}") from e

    def _parse_doc(self, file_path: str) -> str:
        """
        解析旧版 Word 文件 (.doc)

        使用纯 Python 方案（olefile），不依赖 Microsoft Office
        如果可用，也支持 win32com 作为增强方案
        """
        try:
            # 方案 0: 检测是否是 HTML 格式的 .doc 文件（WPS 保存的）
            try:
                with open(file_path, "rb") as f:
                    content = f.read()
                header = content[:200]
                if header.startswith(b"<html") or header.startswith(b"<HTML") or b"<!DOCTYPE" in header:
                    logger_manager.info(f"检测到 HTML 格式的 .doc 文件: {file_path}")
                    return self._extract_text_from_html(content.decode("utf-8", errors="ignore"))
            except Exception as e:
                logger_manager.debug(f"HTML 检测失败: {e}")

            # 方案 1: 使用 olefile 纯 Python 解析（优先，兼容 WPS 用户）
            try:
                import olefile

                if olefile.isOleFile(file_path):
                    text = self._parse_doc_from_ole_robust(file_path)
                    if text and len(text.strip()) > 10:  # 确保提取到有效内容
                        logger_manager.debug(f"成功解析 DOC 文件（通过 olefile）: {file_path}")
                        return text
                    else:
                        logger_manager.warning("olefile 提取的内容过少，尝试其他方案")
                else:
                    logger_manager.warning(f"文件不是有效的 OLE 格式: {file_path}")

            except ImportError:
                logger_manager.info("olefile 未安装，尝试其他方案")
            except Exception as e:
                logger_manager.warning(f"olefile 解析失败: {e}，尝试其他方案")

            # 方案 2: 使用 win32com（需要 Microsoft Office，可选增强）
            try:
                import win32com.client as win32

                logger_manager.info("尝试使用 win32com 解析（需要 Microsoft Office）")
                word = win32.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False  # 不显示警告

                # 转换为临时 DOCX
                doc = word.Documents.Open(os.path.abspath(file_path))
                temp_docx = file_path + ".temp.docx"
                doc.SaveAs2(temp_docx, FileFormat=16)  # 16 = wdFormatXMLDocument
                doc.Close()
                word.Quit()

                # 解析转换后的 DOCX
                from docx import Document

                docx_doc = Document(temp_docx)
                text_parts = [para.text for para in docx_doc.paragraphs if para.text.strip()]

                # 清理临时文件
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)

                logger_manager.debug(f"成功解析 DOC 文件（通过 win32com）: {file_path}")
                return "\n".join(text_parts)

            except ImportError:
                logger_manager.debug("win32com 不可用（正常，WPS 用户不需要）")
            except Exception as e:
                logger_manager.warning(f"win32com 解析失败: {e}")

            # 方案 3: 使用 antiword（Linux/macOS，如果可用）
            try:
                import subprocess

                result = subprocess.run(["antiword", file_path], capture_output=True, text=True, timeout=10)
                if result.returncode == 0 and result.stdout.strip():
                    logger_manager.debug(f"成功解析 DOC 文件（通过 antiword）: {file_path}")
                    return result.stdout
            except FileNotFoundError:
                pass  # antiword 未安装
            except Exception as e:
                logger_manager.debug(f"antiword 不可用: {e}")

            # 如果所有方案都失败，抛出错误
            raise FileCorruptedError(
                "旧版 Word 文件解析失败。请安装 olefile: pip install olefile，或在 WPS 中将文件另存为 .docx 格式"
            )

        except FileCorruptedError:
            raise
        except Exception as e:
            logger_manager.error(f"DOC 解析失败: {file_path}: {e}")
            raise FileCorruptedError(f"旧版 Word 文件解析失败: {e}") from e

    def _parse_html_doc(self, file_path: str) -> str:
        """
        解析 HTML 格式的 .doc 文件（WPS 保存的）

        WPS 有时会将文档保存为 HTML 格式但使用 .doc 扩展名
        """
        try:
            # 读取文件内容
            with open(file_path, "rb") as f:
                raw = f.read()

            # 尝试多种编码
            for encoding in ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]:
                try:
                    html_content = raw.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                html_content = raw.decode("utf-8", errors="ignore")

            return self._extract_text_from_html(html_content)

        except Exception as e:
            logger_manager.error(f"HTML 格式 .doc 解析失败: {e}")
            raise

    def _extract_text_from_html(self, html_content: str) -> str:
        """从 HTML 内容中提取纯文本"""
        import re
        from html.parser import HTMLParser

        class HTMLTextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text_parts = []
                self.skip_tags = {"script", "style", "head"}
                self.current_tag = None

            def handle_starttag(self, tag, attrs):
                self.current_tag = tag.lower()

            def handle_endtag(self, tag):
                if tag.lower() in ("p", "div", "br", "h1", "h2", "h3", "h4", "h5", "h6", "li", "tr"):
                    self.text_parts.append("\n")
                self.current_tag = None

            def handle_data(self, data):
                if self.current_tag not in self.skip_tags:
                    text = data.strip()
                    if text:
                        self.text_parts.append(text)

            def get_text(self):
                return "".join(self.text_parts)

        # 提取文本
        extractor = HTMLTextExtractor()
        extractor.feed(html_content)
        text = extractor.get_text()

        # 清理文本
        text = re.sub(r"\n\s*\n", "\n", text)  # 合并多个空行
        text = re.sub(r"[ \t]+", " ", text)  # 合并多个空格
        text = text.strip()

        if len(text) > 10:
            logger_manager.info(f"成功从 HTML 提取文本，共 {len(text)} 字符")
            return text
        else:
            raise ValueError("HTML 文档内容过少")

    def _parse_doc_bytes(self, file_bytes: bytes) -> str:
        """从字节流解析旧版 Word 文件 (.doc)"""
        import tempfile

        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            return self._parse_doc(tmp_path)
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except OSError as e:
                    logger_manager.warning(f"清理临时文件失败: {e}")

    def _parse_doc_from_ole_robust(self, file_path: str) -> str:
        """
        健壮的 OLE 复合文档文本提取

        从 .doc 文件的多个流中提取文本内容，兼容 WPS 创建的文档
        """
        import re

        import olefile

        try:
            ole = olefile.OleFileIO(file_path)
        except Exception as e:
            logger_manager.error(f"打开 OLE 文件失败: {e}")
            raise

        try:
            # 先检查 WordDocument 流是否包含 HTML 内容（WPS 保存的 .doc 文件）
            try:
                word_stream = ole.openstream("WordDocument").read()
                if word_stream[:100].lstrip().startswith((b"<html", b"<HTML", b"<!DOCTYPE")):
                    logger_manager.info("检测到 OLE 文件中包含 HTML 内容（WPS 格式）")
                    for encoding in ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]:
                        try:
                            html_content = word_stream.decode(encoding)
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        html_content = word_stream.decode("utf-8", errors="ignore")

                    return self._extract_text_from_html(html_content)
            except Exception as e:
                logger_manager.debug(f"检查 WordDocument 流 HTML 内容失败: {e}")

            # 尝试从多个流中提取文本
            text_parts = []

            # 中文标点和特殊字符范围
            # 中文：\u4e00-\u9fa5
            # 中文标点：\u3000-\u303f, \uff00-\uffef
            # 全角字符：\uff01-\uff5e
            chinese_char_pattern = re.compile(
                r"[\u4e00-\u9fa5\u3000-\u303f\uff00-\uffef\u2000-\u206f\u0020-\u007e\r\n\t]"
            )

            # 流名称列表（按优先级排序）
            stream_names = ["WordDocument", "1Table", "0Table", "Data"]

            for stream_name in stream_names:
                if not ole.exists(stream_name):
                    continue

                try:
                    stream = ole.openstream(stream_name)
                    raw_data = stream.read()

                    # 方法 1: UTF-16-LE 解码（Word 文档标准编码）
                    try:
                        decoded = raw_data.decode("utf-16-le", errors="ignore")
                        # 提取中文、英文、标点
                        chars = chinese_char_pattern.findall(decoded)
                        text = "".join(chars)
                        # 清理多余空白
                        text = re.sub(r"\s+", " ", text).strip()
                        if len(text) > 10:  # 降低阈值到 10
                            text_parts.append(text)
                            logger_manager.debug(f"从 {stream_name} 流提取到 {len(text)} 字符")
                    except Exception as e:
                        logger_manager.debug(f"{stream_name} UTF-16-LE 解码失败: {e}")

                    # 方法 2: 逐字节提取（备用方案）
                    if not text_parts:
                        text = ""
                        i = 0
                        while i < len(raw_data) - 1:
                            # ASCII 字符
                            if raw_data[i + 1] == 0 and 32 <= raw_data[i] < 127:
                                text += chr(raw_data[i])
                            # 中文字符范围
                            elif (
                                0x4E00 <= (raw_data[i] | (raw_data[i + 1] << 8)) <= 0x9FA5
                                or 0x3000 <= (raw_data[i] | (raw_data[i + 1] << 8)) <= 0x303F
                                or 0xFF00 <= (raw_data[i] | (raw_data[i + 1] << 8)) <= 0xFFEF
                            ):
                                char_code = raw_data[i] | (raw_data[i + 1] << 8)
                                text += chr(char_code)
                                i += 1
                            i += 1

                        text = re.sub(r"\s+", " ", text).strip()
                        if len(text) > 10:
                            text_parts.append(text)
                            logger_manager.debug(f"从 {stream_name} 流逐字节提取到 {len(text)} 字符")

                except Exception as e:
                    logger_manager.debug(f"读取 {stream_name} 流失败: {e}")
                    continue

            # 合并所有提取到的文本
            if text_parts:
                # 去重（取最长的文本）
                best_text = max(text_parts, key=len)
                logger_manager.info(f"成功解析 DOC 文件，提取到 {len(best_text)} 字符")
                return best_text
            else:
                raise ValueError("无法从 OLE 流中提取文本")

        except Exception as e:
            logger_manager.error(f"OLE 流解析失败: {e}")
            raise
        finally:
            ole.close()

    def _parse_doc_from_ole(self, file_path: str) -> str:
        """
        从 OLE 复合文档中提取文本（备用方案，向后兼容）

        适用于没有安装 pywin32 或 antiword 的情况
        """
        return self._parse_doc_from_ole_robust(file_path)

    def _parse_txt(self, file_path: str) -> str:
        """解析 TXT 文件（支持多编码回退）"""
        for encoding in ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]:
            try:
                with open(file_path, encoding=encoding) as f:
                    text = f.read()
                logger_manager.debug(f"成功解析 TXT 文件（编码: {encoding}）: {file_path}")
                return text
            except UnicodeDecodeError:
                continue
        # 所有编码都失败，使用 replace 模式兜底
        with open(file_path, encoding="utf-8", errors="replace") as f:
            logger_manager.warning(f"TXT 文件编码检测失败，使用 UTF-8 replace 模式: {file_path}")
            return f.read()

    def _parse_txt_bytes(self, file_bytes: bytes) -> str:
        """从字节流解析 TXT（支持多编码回退）"""
        for encoding in ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]:
            try:
                text = file_bytes.decode(encoding)
                logger_manager.debug(f"成功解析 TXT 字节流（编码: {encoding}）")
                return text
            except UnicodeDecodeError:
                continue
        logger_manager.warning("TXT 字节流编码检测失败，使用 UTF-8 replace 模式")
        return file_bytes.decode("utf-8", errors="replace")

    def _split_clauses(self, text: str) -> list[Clause]:
        """
        将文本切分为条款

        基于常见条款标题模式进行切分：
        - 第X条、第X章、第X节
        - 一、二、三、
        - 1.、2.、3.
        - 大写标题（如"违约责任"、"争议解决"）
        - 英文合同条款（Article I, Section 2.1 等）
        """
        import re

        clauses = []
        # 条款标题模式（按优先级排序）
        clause_patterns = [
            # 中文：第一条、第一章
            (r"^(第[一二三四五六七八九十百千\d]+[条章节])\s*", 1),
            # 中文数字序号：一、二、
            (r"^([一二三四五六七八九十]+[、.．])\s*", 2),
            # 中文数字序号：第1条
            (r"^(第\d+条)\s*", 3),
            # 纯数字序号：1.、1、
            (r"^(\d+[.．、])\s*([^\n]{2,30})$", 4),
            # 英文合同：Article I, Article 1
            (r"^(Article\s+[IVXLCDM\d]+)\s*", 5),
            # 英文合同：Section 1.1, Section 2
            (r"^(Section\s+\d+(?:\.\d+)?)\s*", 6),
            # 带括号的序号：（一）、(1)
            (r"^[（(]([一二三四五六七八九十\d]+)[）)]\s*([^\n]{2,30})$", 7),
        ]

        # 常见合同标题关键词（无序号时作为条款标题）
        clause_keywords = [
            "违约责任",
            "争议解决",
            "保密条款",
            "知识产权",
            "不可抗力",
            "生效条款",
            "当事人信息",
            "合同解除",
            "合同终止",
            "付款方式",
            "交货期限",
            "质量保证",
            "售后服务",
            "违约责任",
            "管辖法院",
            "适用法律",
            "定义与解释",
            "权利义务",
            "合作内容",
            "合作期限",
            "费用与支付",
            "保密义务",
            "违约责任",
            "不可抗力",
            "争议解决",
            "其他约定",
            "附则",
            "总则",
        ]

        lines = text.split("\n")
        current_clause = None
        current_content = []
        pos = 0

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if current_content:
                    current_content.append("")
                continue

            is_clause_header = False

            # 策略1：匹配序号模式
            for pattern, _priority in clause_patterns:
                match = re.match(pattern, stripped)
                if match:
                    # 保存之前的条款
                    if current_clause is not None:
                        content = "\n".join(current_content).strip()
                        if content:
                            clauses.append(
                                Clause(
                                    id=len(clauses) + 1,
                                    content=content,
                                    title=current_clause,
                                    start_pos=pos - len(content),
                                    end_pos=pos,
                                )
                            )

                    current_clause = stripped
                    current_content = []
                    is_clause_header = True
                    break

            # 策略2：无序号时，匹配常见条款标题关键词
            if not is_clause_header and len(stripped) <= 30:
                for keyword in clause_keywords:
                    if stripped == keyword or stripped.startswith(keyword + "：") or stripped.startswith(keyword + " "):
                        if current_clause is not None:
                            content = "\n".join(current_content).strip()
                            if content:
                                clauses.append(
                                    Clause(
                                        id=len(clauses) + 1,
                                        content=content,
                                        title=current_clause,
                                        start_pos=pos - len(content),
                                        end_pos=pos,
                                    )
                                )

                        current_clause = stripped
                        current_content = []
                        is_clause_header = True
                        break

            if not is_clause_header:
                current_content.append(stripped)

            pos += len(line) + 1  # +1 for newline

        # 保存最后一个条款
        if current_content:
            content = "\n".join(current_content).strip()
            if content:
                clauses.append(
                    Clause(
                        id=len(clauses) + 1,
                        content=content,
                        title=current_clause,
                        start_pos=pos - len(content),
                        end_pos=pos,
                    )
                )

        # 如果没有检测到任何条款结构，尝试按段落切分
        if not clauses and text.strip():
            paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
            if len(paragraphs) > 1:
                for i, para in enumerate(paragraphs):
                    clauses.append(
                        Clause(
                            id=i + 1,
                            content=para,
                            title=None,
                            start_pos=text.find(para),
                            end_pos=text.find(para) + len(para),
                        )
                    )
            else:
                # 整个文档作为一个条款
                clauses.append(Clause(id=1, content=text.strip(), title=None, start_pos=0, end_pos=len(text)))

        return clauses

    def detect_clause_type(self, clause: Clause) -> str | None:
        """
        检测条款类型

        Args:
            clause: 条款对象

        Returns:
            条款类型字符串
        """
        type_keywords = {
            "违约责任": ["违约责任", "违约", "违约金", "赔偿", "损害赔偿"],
            "争议解决": ["争议解决", "争议", "仲裁", "诉讼", "管辖", "法院"],
            "保密条款": ["保密", "商业秘密", "机密", "confidential"],
            "知识产权": ["知识产权", "版权", "专利", "商标", "著作权"],
            "解除条款": ["解除", "终止", "合同终止", "合同解除"],
            "不可抗力": ["不可抗力", "force majeure"],
            "生效条款": ["生效", "本合同自", "签字盖章"],
            "当事人信息": ["甲方", "乙方", "当事人", "地址", "法定代表人"],
        }

        text = (clause.title or "") + " " + clause.content[:500]

        for clause_type, keywords in type_keywords.items():
            for keyword in keywords:
                if keyword.lower() in text.lower():
                    return clause_type

        return None
