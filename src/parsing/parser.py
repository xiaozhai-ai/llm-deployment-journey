"""
文件解析模块
支持 PDF、Word (.docx)、TXT 格式文件的解析与结构化
增强版：统一异常处理 + 日志记录 + OCR 扫描件支持
"""

from __future__ import annotations

import os
import sys
from dataclasses import dataclass, field

from src.core.exceptions import FileCorruptedError, FileTooLargeError, UnsupportedFormatError
from src.infra.logger import logger_manager
from src.parsing.scan_detector import ScanDetector


@dataclass
class Clause:
    """条款数据结构"""

    id: int
    content: str
    title: str | None = None
    start_pos: int = 0
    end_pos: int = 0
    clause_type: str | None = None  # 条款类型（如违约责任、争议解决等）


@dataclass
class ParsedDocument:
    """解析后的文档结构"""

    original_filename: str
    file_type: str  # pdf, docx, doc, txt
    full_text: str
    clauses: list[Clause]
    metadata: dict | None = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class EnhancedParsedDocument(ParsedDocument):
    """增强版解析结果（向后兼容 ParsedDocument，新增版面/结构/法律要素字段）"""

    # 版面分析（OCR 后填充）
    pages: list = field(default_factory=list)  # list[PageLayout]
    # 条款层级树
    clause_tree: list = field(default_factory=list)  # list[ClauseNode]
    # 法律要素
    legal_metadata: object | None = None  # LegalMetadata
    amounts: list = field(default_factory=list)  # list[MoneyAmount]
    dates: list = field(default_factory=list)  # list[DateEntity]
    signatures: list = field(default_factory=list)  # list[SignatureInfo]
    definitions: list = field(default_factory=list)  # list[Definition]
    revisions: list = field(default_factory=list)  # list[Revision]
    # 引擎信息
    ocr_used: bool = False
    scan_detection: object | None = None  # ScanDetectionResult
    engine_info: dict = field(default_factory=dict)


class DocumentParser:
    """文档解析器"""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}

    def __init__(self, max_file_size_mb: int = 10, ocr_enabled: bool = False):
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        self.ocr_enabled = ocr_enabled
        self._scan_detector = None
        self._layout_engine = None

    def parse_file(self, file_path: str) -> EnhancedParsedDocument:
        """
        解析文件，返回增强版结构化文档

        Args:
            file_path: 文件路径

        Returns:
            EnhancedParsedDocument: 增强版解析结果（兼容 ParsedDocument）

        Raises:
            UnsupportedFormatError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size_bytes:
            raise FileTooLargeError(file_size / 1024 / 1024, self.max_file_size_bytes / 1024 / 1024)

        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedFormatError(ext)

        # PDF 扫描件检测 + OCR 路由
        if ext == ".pdf" and self.ocr_enabled:
            ocr_result = self._try_ocr_pdf_from_path(file_path)
            if ocr_result is not None:
                return ocr_result

        if ext == ".pdf":
            full_text = self._parse_pdf(file_path)
        elif ext == ".docx":
            full_text = self._parse_docx(file_path)
        elif ext == ".doc":
            full_text = self._parse_doc(file_path)
        elif ext == ".txt":
            full_text = self._parse_txt(file_path)
        else:
            raise UnsupportedFormatError(ext)

        clauses = self._split_clauses(full_text)

        return EnhancedParsedDocument(
            original_filename=os.path.basename(file_path),
            file_type=ext[1:],
            full_text=full_text,
            clauses=clauses,
            metadata={"file_size": os.path.getsize(file_path), "extension": ext},
        )

    def parse_bytes(self, file_bytes: bytes, filename: str) -> EnhancedParsedDocument:
        """
        从字节流解析文件（用于 Web 上传）

        Args:
            file_bytes: 文件字节流
            filename: 文件名

        Returns:
            EnhancedParsedDocument: 增强版解析结果（兼容 ParsedDocument）
        """
        if len(file_bytes) > self.max_file_size_bytes:
            raise FileTooLargeError(len(file_bytes) / 1024 / 1024, self.max_file_size_bytes / 1024 / 1024)

        ext = os.path.splitext(filename)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedFormatError(ext)

        # PDF 扫描件检测 + OCR 路由
        if ext == ".pdf" and self.ocr_enabled:
            ocr_result = self._try_ocr_pdf_bytes(file_bytes, filename)
            if ocr_result is not None:
                return ocr_result

        if ext == ".pdf":
            full_text = self._parse_pdf_bytes(file_bytes)
        elif ext == ".docx":
            full_text = self._parse_docx_bytes(file_bytes)
        elif ext == ".doc":
            full_text = self._parse_doc_bytes(file_bytes)
        elif ext == ".txt":
            full_text = self._parse_txt_bytes(file_bytes)
        else:
            raise UnsupportedFormatError(ext)

        clauses = self._split_clauses(full_text)

        return EnhancedParsedDocument(
            original_filename=filename,
            file_type=ext[1:],
            full_text=full_text,
            clauses=clauses,
            metadata={"file_size": len(file_bytes), "extension": ext},
        )

    def _parse_pdf(self, file_path: str) -> str:
        """解析 PDF 文件（含表格内容）"""
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                return self._extract_pdf_text(pdf, source=file_path)
        except FileCorruptedError:
            raise
        except Exception as e:
            logger_manager.error(f"PDF 解析失败: {file_path}: {e}")
            raise FileCorruptedError(f"PDF 文件解析失败: {e}") from e

    def _parse_pdf_bytes(self, file_bytes: bytes) -> str:
        """从字节流解析 PDF（含表格内容）"""
        try:
            import io

            import pdfplumber

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                return self._extract_pdf_text(pdf, source="bytes")
        except FileCorruptedError:
            raise
        except Exception as e:
            logger_manager.error(f"PDF 字节流解析失败: {e}")
            raise FileCorruptedError(f"PDF 字节流解析失败: {e}") from e

    @staticmethod
    def _extract_pdf_text(pdf, source: str = "") -> str:
        """从已打开的 pdfplumber PDF 对象提取文本和表格，per-page 容错。"""
        text_parts: list[str] = []
        failed_pages = 0
        for page_idx, page in enumerate(pdf.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        cells = [str(c).strip() if c else "" for c in row]
                        row_text = " | ".join(c for c in cells if c)
                        if row_text:
                            text_parts.append(row_text)
            except Exception as e:
                failed_pages += 1
                logger_manager.warning(f"PDF 第 {page_idx + 1} 页解析失败，已跳过: {e}")

        if not text_parts:
            raise FileCorruptedError(f"PDF 所有页面均解析失败{f' ({source})' if source else ''}")
        if failed_pages:
            logger_manager.warning(f"PDF 共 {failed_pages} 页解析失败，已跳过")

        logger_manager.debug(f"成功解析 PDF{f' ({source})' if source else ''}，共 {len(text_parts)} 段")
        return "\n".join(text_parts)

    def _parse_docx(self, file_path: str) -> str:
        """解析 Word 文件（含表格内容）"""
        try:
            from docx import Document

            doc = Document(file_path)
            text = self._extract_docx_text(doc)
            logger_manager.debug(f"成功解析 DOCX 文件: {file_path}")
            return text
        except Exception as e:
            logger_manager.error(f"DOCX 解析失败: {file_path}: {e}")
            raise FileCorruptedError(f"Word 文件解析失败: {e}") from e

    def _parse_docx_bytes(self, file_bytes: bytes) -> str:
        """从字节流解析 Word（含表格内容）"""
        try:
            import io

            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            text = self._extract_docx_text(doc)
            logger_manager.debug("成功解析 DOCX 字节流")
            return text
        except Exception as e:
            logger_manager.error(f"DOCX 字节流解析失败: {e}")
            raise FileCorruptedError(f"Word 字节流解析失败: {e}") from e

    @staticmethod
    def _extract_docx_text(doc) -> str:
        """从已打开的 python-docx Document 对象提取文本和表格（按文档顺序）。"""
        # 预构建 element → object 映射，避免 O(n²) 查找
        para_map = {p._element: p for p in doc.paragraphs}
        table_map = {t._element: t for t in doc.tables}

        text_parts: list[str] = []
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag == "p":
                para = para_map.get(element)
                if para and para.text.strip():
                    text_parts.append(para.text)
            elif tag == "tbl":
                table = table_map.get(element)
                if table:
                    table_lines = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        row_text = " | ".join(c for c in cells if c)
                        if row_text:
                            table_lines.append(row_text)
                    if table_lines:
                        text_parts.append("\n".join(table_lines))

        return "\n".join(text_parts)

    def _parse_doc(self, file_path: str) -> str:
        """
        解析旧版 Word 文件 (.doc)

        使用纯 Python 方案（olefile），不依赖 Microsoft Office
        如果可用，也支持 win32com 作为增强方案
        """
        try:
            # 方案 0: 检测是否是 HTML 格式的 .doc 文件（WPS 保存的）
            try:
                with open(file_path, "rb") as f:
                    content = f.read()
                header = content[:200]
                if header.startswith(b"<html") or header.startswith(b"<HTML") or b"<!DOCTYPE" in header:
                    logger_manager.info(f"检测到 HTML 格式的 .doc 文件: {file_path}")
                    return self._extract_text_from_html(content.decode("utf-8", errors="ignore"))
            except Exception as e:
                logger_manager.debug(f"HTML 检测失败: {e}")

            # 方案 1: 使用 olefile 纯 Python 解析（优先，兼容 WPS 用户）
            try:
                import olefile

                if olefile.isOleFile(file_path):
                    text = self._parse_doc_from_ole_robust(file_path)
                    if text and len(text.strip()) > 3:  # 宽松阈值，避免丢弃短文档
                        logger_manager.debug(f"成功解析 DOC 文件（通过 olefile）: {file_path}")
                        return text
                    else:
                        logger_manager.warning("olefile 提取的内容过少，尝试其他方案")
                else:
                    logger_manager.warning(f"文件不是有效的 OLE 格式: {file_path}")

            except ImportError:
                logger_manager.info("olefile 未安装，尝试其他方案")
            except Exception as e:
                logger_manager.warning(f"olefile 解析失败: {e}，尝试其他方案")

            # 方案 2: 使用 win32com（需要 Microsoft Office，可选增强）
            try:
                import win32com.client as win32

                logger_manager.info("尝试使用 win32com 解析（需要 Microsoft Office）")
                word = win32.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False  # 不显示警告

                try:
                    # 转换为临时 DOCX
                    doc = word.Documents.Open(os.path.abspath(file_path))
                    temp_docx = file_path + ".temp.docx"
                    doc.SaveAs2(temp_docx, FileFormat=16)  # 16 = wdFormatXMLDocument
                    doc.Close()

                    # 解析转换后的 DOCX
                    from docx import Document

                    docx_doc = Document(temp_docx)
                    text = self._extract_docx_text(docx_doc)

                    # 清理临时文件
                    if os.path.exists(temp_docx):
                        os.remove(temp_docx)
                finally:
                    try:
                        word.Quit()
                    except Exception:
                        pass  # 确保 Word 进程被释放

                logger_manager.debug(f"成功解析 DOC 文件（通过 win32com）: {file_path}")
                return text

            except ImportError:
                logger_manager.debug("win32com 不可用（正常，WPS 用户不需要）")
            except Exception as e:
                logger_manager.warning(f"win32com 解析失败: {e}")

            # 方案 3: 使用 antiword（仅 Linux/macOS）
            if sys.platform != "win32":
                try:
                    import subprocess

                    result = subprocess.run(["antiword", file_path], capture_output=True, text=True, timeout=10)
                    if result.returncode == 0 and result.stdout.strip():
                        logger_manager.debug(f"成功解析 DOC 文件（通过 antiword）: {file_path}")
                        return result.stdout
                except FileNotFoundError:
                    pass  # antiword 未安装
                except Exception as e:
                    logger_manager.debug(f"antiword 不可用: {e}")

            # 如果所有方案都失败，抛出错误
            raise FileCorruptedError(
                "旧版 Word 文件解析失败。请安装 olefile: pip install olefile，或在 WPS 中将文件另存为 .docx 格式"
            )

        except FileCorruptedError:
            raise
        except Exception as e:
            logger_manager.error(f"DOC 解析失败: {file_path}: {e}")
            raise FileCorruptedError(f"旧版 Word 文件解析失败: {e}") from e

    def _extract_text_from_html(self, html_content: str) -> str:
        """从 HTML 内容中提取纯文本"""
        import re
        from html.parser import HTMLParser

        class HTMLTextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text_parts = []
                self.skip_tags = {"script", "style", "head"}
                self.current_tag = None

            def handle_starttag(self, tag, attrs):
                self.current_tag = tag.lower()

            def handle_endtag(self, tag):
                if tag.lower() in ("p", "div", "br", "h1", "h2", "h3", "h4", "h5", "h6", "li", "tr"):
                    self.text_parts.append("\n")
                self.current_tag = None

            def handle_data(self, data):
                if self.current_tag not in self.skip_tags:
                    text = data.strip()
                    if text:
                        self.text_parts.append(text)

            def get_text(self):
                return "".join(self.text_parts)

        # 提取文本
        extractor = HTMLTextExtractor()
        extractor.feed(html_content)
        text = extractor.get_text()

        # 清理文本
        text = re.sub(r"\n\s*\n", "\n", text)  # 合并多个空行
        text = re.sub(r"[ \t]+", " ", text)  # 合并多个空格
        text = text.strip()

        if len(text) > 10:
            logger_manager.info(f"成功从 HTML 提取文本，共 {len(text)} 字符")
            return text
        else:
            raise ValueError("HTML 文档内容过少")

    def _parse_doc_bytes(self, file_bytes: bytes) -> str:
        """从字节流解析旧版 Word 文件 (.doc)"""
        import tempfile

        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            return self._parse_doc(tmp_path)
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except OSError as e:
                    logger_manager.warning(f"清理临时文件失败: {e}")

    def _parse_doc_from_ole_robust(self, file_path: str) -> str:
        """
        健壮的 OLE 复合文档文本提取

        从 .doc 文件的多个流中提取文本内容，兼容 WPS 创建的文档
        """
        import re

        import olefile

        try:
            ole = olefile.OleFileIO(file_path)
        except Exception as e:
            logger_manager.error(f"打开 OLE 文件失败: {e}")
            raise

        try:
            # 先检查 WordDocument 流是否包含 HTML 内容（WPS 保存的 .doc 文件）
            try:
                word_stream = ole.openstream("WordDocument").read()
                if word_stream[:100].lstrip().startswith((b"<html", b"<HTML", b"<!DOCTYPE")):
                    logger_manager.info("检测到 OLE 文件中包含 HTML 内容（WPS 格式）")
                    for encoding in ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]:
                        try:
                            html_content = word_stream.decode(encoding)
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        html_content = word_stream.decode("utf-8", errors="ignore")

                    return self._extract_text_from_html(html_content)
            except Exception as e:
                logger_manager.debug(f"检查 WordDocument 流 HTML 内容失败: {e}")

            # 尝试从多个流中提取文本
            text_parts = []

            # 中文标点和特殊字符范围
            # 中文：\u4e00-\u9fa5
            # 中文标点：\u3000-\u303f, \uff00-\uffef
            # 全角字符：\uff01-\uff5e
            chinese_char_pattern = re.compile(
                r"[\u4e00-\u9fa5\u3000-\u303f\uff00-\uffef\u2000-\u206f\u0020-\u007e\r\n\t]"
            )

            # 流名称列表（按优先级排序）
            stream_names = ["WordDocument", "1Table", "0Table", "Data"]

            for stream_name in stream_names:
                if not ole.exists(stream_name):
                    continue

                try:
                    stream = ole.openstream(stream_name)
                    raw_data = stream.read()

                    # 方法 1: UTF-16-LE 解码（Word 文档标准编码）
                    try:
                        decoded = raw_data.decode("utf-16-le", errors="ignore")
                        # 提取中文、英文、标点
                        chars = chinese_char_pattern.findall(decoded)
                        text = "".join(chars)
                        # 清理多余空白
                        text = re.sub(r"\s+", " ", text).strip()
                        if len(text) > 3:
                            text_parts.append(text)
                            logger_manager.debug(f"从 {stream_name} 流提取到 {len(text)} 字符")
                    except Exception as e:
                        logger_manager.debug(f"{stream_name} UTF-16-LE 解码失败: {e}")

                    # 方法 2: 逐字节提取（备用方案）
                    if not text_parts:
                        text = ""
                        i = 0
                        while i < len(raw_data) - 1:
                            # ASCII 字符
                            if raw_data[i + 1] == 0 and 32 <= raw_data[i] < 127:
                                text += chr(raw_data[i])
                            # 中文字符范围
                            elif (
                                0x4E00 <= (raw_data[i] | (raw_data[i + 1] << 8)) <= 0x9FA5
                                or 0x3000 <= (raw_data[i] | (raw_data[i + 1] << 8)) <= 0x303F
                                or 0xFF00 <= (raw_data[i] | (raw_data[i + 1] << 8)) <= 0xFFEF
                            ):
                                char_code = raw_data[i] | (raw_data[i + 1] << 8)
                                text += chr(char_code)
                                i += 1
                            i += 1

                        text = re.sub(r"\s+", " ", text).strip()
                        if len(text) > 3:
                            text_parts.append(text)
                            logger_manager.debug(f"从 {stream_name} 流逐字节提取到 {len(text)} 字符")

                except Exception as e:
                    logger_manager.debug(f"读取 {stream_name} 流失败: {e}")
                    continue

            # 合并所有提取到的文本
            if text_parts:
                # 去重（取最长的文本）
                best_text = max(text_parts, key=len)
                logger_manager.info(f"成功解析 DOC 文件，提取到 {len(best_text)} 字符")
                return best_text
            else:
                raise ValueError("无法从 OLE 流中提取文本")

        except Exception as e:
            logger_manager.error(f"OLE 流解析失败: {e}")
            raise
        finally:
            ole.close()

    def _parse_txt(self, file_path: str) -> str:
        """解析 TXT 文件（支持多编码回退，自动处理平台换行符）"""
        for encoding in self._ENCODING_FALLBACK_CHAIN:
            try:
                with open(file_path, encoding=encoding) as f:
                    text = f.read()
                logger_manager.debug(f"成功解析 TXT 文件: {file_path}")
                return text
            except UnicodeDecodeError:
                continue
        # 所有编码都失败，使用 replace 模式兜底
        with open(file_path, encoding="utf-8", errors="replace") as f:
            logger_manager.warning(f"TXT 文件编码检测失败，使用 UTF-8 replace 模式: {file_path}")
            return f.read()

    def _parse_txt_bytes(self, file_bytes: bytes) -> str:
        """从字节流解析 TXT（支持多编码回退）"""
        text = self._decode_text_bytes(file_bytes)
        logger_manager.debug("成功解析 TXT 字节流")
        return text

    _ENCODING_FALLBACK_CHAIN = ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]

    @staticmethod
    def _decode_text_bytes(raw: bytes) -> str:
        """尝试多种编码解码字节流，全部失败则用 utf-8 replace 兜底。"""
        for encoding in DocumentParser._ENCODING_FALLBACK_CHAIN:
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        logger_manager.warning("TXT 编码检测失败，使用 UTF-8 replace 模式")
        return raw.decode("utf-8", errors="replace")

    @staticmethod
    def _split_clauses(text: str) -> list[Clause]:
        """
        将文本切分为条款

        基于常见条款标题模式进行切分：
        - 第X条、第X章、第X节
        - 一、二、三、
        - 1.、2.、3.
        - 大写标题（如"违约责任"、"争议解决"）
        - 英文合同条款（Article I, Section 2.1 等）
        """
        from src.parsing.structure.clause_patterns import CLAUSE_PATTERNS, CLAUSE_TITLE_KEYWORDS

        clauses: list[Clause] = []

        clause_patterns = CLAUSE_PATTERNS  # [(priority, compiled_pattern, desc), ...]
        clause_keywords = CLAUSE_TITLE_KEYWORDS

        lines = text.split("\n")
        current_clause = None
        current_content = []
        pos = 0

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if current_content:
                    current_content.append("")
                continue

            is_clause_header = False

            # 策略1：匹配序号模式
            for _priority, pattern, _desc in clause_patterns:
                match = pattern.match(stripped)
                if match:
                    # 保存之前的条款
                    if current_clause is not None:
                        content = "\n".join(current_content).strip()
                        if content:
                            clauses.append(
                                Clause(
                                    id=len(clauses) + 1,
                                    content=content,
                                    title=current_clause,
                                    start_pos=pos - len(content),
                                    end_pos=pos,
                                )
                            )

                    current_clause = stripped
                    current_content = []
                    is_clause_header = True
                    break

            # 策略2：无序号时，匹配常见条款标题关键词
            if not is_clause_header and len(stripped) <= 30:
                for keyword in clause_keywords:
                    if stripped == keyword or stripped.startswith(keyword + "：") or stripped.startswith(keyword + " "):
                        if current_clause is not None:
                            content = "\n".join(current_content).strip()
                            if content:
                                clauses.append(
                                    Clause(
                                        id=len(clauses) + 1,
                                        content=content,
                                        title=current_clause,
                                        start_pos=pos - len(content),
                                        end_pos=pos,
                                    )
                                )

                        current_clause = stripped
                        current_content = []
                        is_clause_header = True
                        break

            if not is_clause_header:
                current_content.append(stripped)

            pos += len(line) + 1  # +1 for newline

        # 保存最后一个条款
        if current_content:
            content = "\n".join(current_content).strip()
            if content:
                clauses.append(
                    Clause(
                        id=len(clauses) + 1,
                        content=content,
                        title=current_clause,
                        start_pos=pos - len(content),
                        end_pos=pos,
                    )
                )

        # 如果没有检测到任何条款结构，尝试按段落切分
        if not clauses and text.strip():
            paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
            if len(paragraphs) > 1:
                for i, para in enumerate(paragraphs):
                    clauses.append(
                        Clause(
                            id=i + 1,
                            content=para,
                            title=None,
                            start_pos=text.find(para),
                            end_pos=text.find(para) + len(para),
                        )
                    )
            else:
                # 整个文档作为一个条款
                clauses.append(Clause(id=1, content=text.strip(), title=None, start_pos=0, end_pos=len(text)))

        return clauses

    @staticmethod
    def detect_clause_type(clause: Clause) -> str | None:
        """
        检测条款类型

        TODO: 集成到 build_clause_tree / _build_enhanced_from_ocr 流程中，
              当前未被调用。需要统一 Clause/LayoutBlock 接口后接入。

        Args:
            clause: 条款对象（需有 title 和 content 属性）

        Returns:
            条款类型字符串
        """
        type_keywords = {
            "违约责任": ["违约责任", "违约", "违约金", "赔偿", "损害赔偿"],
            "争议解决": ["争议解决", "争议", "仲裁", "诉讼", "管辖", "法院"],
            "保密条款": ["保密", "商业秘密", "机密", "confidential"],
            "知识产权": ["知识产权", "版权", "专利", "商标", "著作权"],
            "解除条款": ["解除", "终止", "合同终止", "合同解除"],
            "不可抗力": ["不可抗力", "force majeure"],
            "生效条款": ["生效", "本合同自", "签字盖章"],
            "当事人信息": ["甲方", "乙方", "当事人", "地址", "法定代表人"],
        }

        text = (clause.title or "") + " " + clause.content[:500]

        for clause_type, keywords in type_keywords.items():
            for keyword in keywords:
                if keyword.lower() in text.lower():
                    return clause_type

        return None

    # ----------------------------------------------------------------
    # OCR 路由（扫描件检测 + 版面分析引擎调用 + fallback）
    # ----------------------------------------------------------------

    def _get_scan_detector(self) -> ScanDetector:
        """懒加载扫描件检测器"""
        if self._scan_detector is None:
            self._scan_detector = ScanDetector()
        return self._scan_detector

    def _get_layout_engine(self):
        """懒加载版面分析引擎，不可用时返回 None"""
        if self._layout_engine is None:
            try:
                from src.parsing.layout.engine import get_layout_engine

                self._layout_engine = get_layout_engine()
            except Exception as e:
                logger_manager.warning(f"版面分析引擎加载失败: {e}")
                self._layout_engine = None
        return self._layout_engine

    def _try_ocr_pdf_from_path(self, file_path: str) -> EnhancedParsedDocument | None:
        """
        尝试 OCR 解析 PDF 文件（路径模式）

        Returns:
            EnhancedParsedDocument（OCR 成功时）或 None（回退到 pdfplumber）
        """
        try:
            detector = self._get_scan_detector()
            scan_result = detector.detect_from_path(file_path)
            logger_manager.info(
                f"扫描件检测: is_scanned={scan_result.is_scanned}, "
                f"density={scan_result.text_density:.3f}, reason={scan_result.reason}"
            )

            if not scan_result.is_scanned:
                return None  # 电子原生 → 回退 pdfplumber

            engine = self._get_layout_engine()
            if engine is None:
                logger_manager.warning("版面分析引擎不可用，回退到 pdfplumber")
                return None

            with open(file_path, "rb") as f:
                pdf_bytes = f.read()

            pages = engine.analyze_pdf(pdf_bytes)
            return self._build_enhanced_from_ocr(
                pages=pages,
                filename=os.path.basename(file_path),
                scan_result=scan_result,
                engine_name=engine.engine_name,
                file_size=os.path.getsize(file_path),
            )
        except Exception as e:
            logger_manager.warning(f"OCR 路径解析失败，回退到 pdfplumber: {e}")
            return None

    def _try_ocr_pdf_bytes(self, file_bytes: bytes, filename: str) -> EnhancedParsedDocument | None:
        """
        尝试 OCR 解析 PDF 字节流（Web 上传模式）

        Returns:
            EnhancedParsedDocument（OCR 成功时）或 None（回退到 pdfplumber）
        """
        try:
            detector = self._get_scan_detector()
            scan_result = detector.detect_from_bytes(file_bytes)
            logger_manager.info(
                f"扫描件检测: is_scanned={scan_result.is_scanned}, "
                f"density={scan_result.text_density:.3f}, reason={scan_result.reason}"
            )

            if not scan_result.is_scanned:
                return None

            engine = self._get_layout_engine()
            if engine is None:
                logger_manager.warning("版面分析引擎不可用，回退到 pdfplumber")
                return None

            pages = engine.analyze_pdf(file_bytes)
            return self._build_enhanced_from_ocr(
                pages=pages,
                filename=filename,
                scan_result=scan_result,
                engine_name=engine.engine_name,
                file_size=len(file_bytes),
            )
        except Exception as e:
            logger_manager.warning(f"OCR 字节流解析失败，回退到 pdfplumber: {e}")
            return None

    @staticmethod
    def _build_enhanced_from_ocr(
        pages: list,
        filename: str,
        scan_result: object,
        engine_name: str,
        file_size: int,
    ) -> EnhancedParsedDocument:
        """
        将 OCR 版面分析结果组装为 EnhancedParsedDocument

        按阅读顺序拼接全文 → 分割条款 → 填充增强字段
        """
        # 按页面阅读顺序拼接全文
        text_parts: list[str] = []
        for page in pages:
            ordered_blocks = page.get_reading_ordered_blocks()
            for block in ordered_blocks:
                if block.block_type in ("text", "title") and block.content.strip():
                    text_parts.append(block.content.strip())
                elif block.block_type == "table" and block.content.strip():
                    # 表格内容（HTML 或文本）也纳入全文
                    text_parts.append(block.content.strip())

        full_text = "\n".join(text_parts)

        # 条款分割（静态方法，无需实例化）
        clauses = DocumentParser._split_clauses(full_text)

        # 构建条款层级树（从 LayoutBlock 保留 bbox/page 信息）
        all_blocks = []
        for page in pages:
            all_blocks.extend(page.get_reading_ordered_blocks())
        try:
            from src.parsing.structure.clause_tree import build_clause_tree_from_blocks

            clause_tree = build_clause_tree_from_blocks(all_blocks)
        except Exception as e:
            logger_manager.warning(f"条款层级树构建失败: {e}")
            clause_tree = []

        # 法律实体提取
        legal_metadata = None
        amounts: list = []
        dates: list = []
        signatures: list = []
        definitions: list = []
        try:
            from src.parsing.legal_entities.amount import extract_amounts
            from src.parsing.legal_entities.date_extractor import extract_dates
            from src.parsing.legal_entities.definition import extract_definitions
            from src.parsing.legal_entities.metadata import extract_metadata
            from src.parsing.legal_entities.signature import detect_signatures

            legal_metadata = extract_metadata(full_text)
            amounts = extract_amounts(full_text)
            dates = extract_dates(full_text)
            signatures = detect_signatures(full_text, clauses=clauses, pages=pages)
            definitions = extract_definitions(full_text, clauses=clauses)
            logger_manager.info(
                f"法律实体提取完成: 金额 {len(amounts)} 个, 日期 {len(dates)} 个, "
                f"签章 {len(signatures)} 个, 定义 {len(definitions)} 个"
            )
        except Exception as e:
            logger_manager.warning(f"法律实体提取失败: {e}")

        return EnhancedParsedDocument(
            original_filename=filename,
            file_type="pdf",
            full_text=full_text,
            clauses=clauses,
            metadata={"file_size": file_size, "extension": ".pdf"},
            # 增强字段
            pages=pages,
            clause_tree=clause_tree,
            legal_metadata=legal_metadata,
            amounts=amounts,
            dates=dates,
            signatures=signatures,
            definitions=definitions,
            ocr_used=True,
            scan_detection=scan_result,
            engine_info={"engine": engine_name, "page_count": len(pages)},
        )
